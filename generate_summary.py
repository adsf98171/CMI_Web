# tasks/generate_summary.py
from flask import Blueprint, request, jsonify, send_file, session
from common.azure_client import get_client_and_deployment
from common.utils import read_word_file, read_pdf_file  # 你的檔案讀取函數
from werkzeug.utils import secure_filename
from docx import Document  # 新增：處理 Word
import os
import json
import tempfile
import uuid  # 用來產生唯一檔名
from docx.shared import Pt
from docx.oxml.ns import qn # 字體
from docx.enum.text import WD_LINE_SPACING # 行距

summary_bp = Blueprint('summary', __name__)

# ============ 多種 Summary 模板 Prompt ============
SUMMARY_TEMPLATES = {
    "quality_control": """[僅使用此模板時啟動]
你是一位在醫學中心服務的品質管理顧問及經營分析專家，請根據提供的數據及報告產生一份嚴謹的經營分析摘要。

重點包含：
1. 年度趨勢比較：分析三年分數變化，找出持續改善或退步的面向。
2. 部門差異：比較不同科別或單位的分數，凸顯強項與弱項。
3. 流程改善建議：針對低分項目提出流程優化建議（如候診時間、資訊透明度）。
4. 後續行動計畫：列出短期、中期、長期的改善措施與追蹤方式。

輸出格式：
**歷年趨勢(跟前幾年比較)** ...，**同儕比較(各病房/各科別)**
**改善建議** ...，**短/中/長期計畫**""",

    "family_medicine": """[僅使用此模板時啟動]
你是一位資深家醫科醫師，請根據病歷產生一份溫暖、易懂的家醫摘要，適合轉診或衛教使用。

重點包含：
1. 病人故事（主訴、病程）
2. 目前診斷與治療
3. 出院後注意事項
4. 生活建議與追蹤計畫

輸出格式：
【家醫科摘要】
親愛的 {病人姓名}：
您這次住院的主要原因是...
目前診斷為...
治療經過...
出院後請注意...
下次門診時間...
祝您早日康復！
""",

    "general": """[僅使用此模板時啟動]
你是一位專業醫療摘要專家，請產生一份簡潔的通用摘要。

重點：
- 主訴
- 主要診斷
- 重要檢查發現
- 治療經過
- 出院狀況與建議

輸出格式：
**摘要**
...
"""
}

# 儲存 Prompt（上限 20 個）
# 1. 指定您想要的資料夾路徑
SAVE_DIR = r"C:\Users\482525\Prompt_File"
# 2. 指定完整檔案路徑
PROMPT_FILE = os.path.join(SAVE_DIR, 'saved_prompts.json')

def load_from_file():
    # 如果檔案夾或檔案不存在，回傳空列表
    if os.path.exists(PROMPT_FILE):
        try:
            with open(PROMPT_FILE, 'r', encoding='utf-8') as f:
                return json.load(f)
        except Exception as e:
            print(f"讀取錯誤: {e}")
            return []
    return []

def save_to_file(prompts):
    # 自動檢查資料夾是否存在，不存在就建立它
    if not os.path.exists(SAVE_DIR):
        os.makedirs(SAVE_DIR)
        
    with open(PROMPT_FILE, 'w', encoding='utf-8') as f:
        # ensure_ascii=False 確保中文字不會變成編碼，indent=4 讓格式好看
        json.dump(prompts, f, ensure_ascii=False, indent=4)

@summary_bp.route('/save_prompt', methods=['POST'])
def save_prompt():
    data = request.json
    name = data.get('name', '').strip()
    prompt = data.get('prompt', '').strip()
    
    if not name or not prompt:
        return jsonify({"error": "名稱與 Prompt 必填"}), 400
    
    # 從「檔案」讀取目前的列表
    prompts = load_from_file()
    
    # 避免重複名稱（刪除舊的同名項）
    prompts = [p for p in prompts if p['name'] != name]
    
    # 加入新資料
    prompts.append({'name': name, 'prompt': prompt})
    
    # 上限 10 個
    if len(prompts) > 10:
        prompts = prompts[-10:]
    
    # 儲存回「檔案」
    save_to_file(prompts)
    
    return jsonify({"success": True})

@summary_bp.route('/load_saved_prompts')
def load_saved_prompts():
    # 這裡必須調用讀取 C:\Users\482525\Prompt_File 的邏輯
    prompts = load_from_file() 
    print(f"DEBUG: 撈取到的資料為 {prompts}") # 加這行在終端機檢查
    return jsonify({"prompts": prompts})

# 新增：插入摘要到 Word 書籤位置
def insert_summary_at_bookmark(word_path, summary_text, bookmark_name="AI_SUMMARY_HERE", output_path=None):
    doc = Document(word_path)
    
    bookmark_found = False
    
    for paragraph in doc.paragraphs:
        if bookmark_name in paragraph._element.xml:
            paragraph.clear()
            
            # === 標題 ===
            title_p = paragraph.insert_paragraph_before("【產生報告總結及建議】")
            title_p.paragraph_format.line_spacing = 1.5
            title_run = title_p.runs[0]
            title_run.bold = True
            title_run.font.size = Pt(16)
            
            # 強制三種字體都設為標楷體（標題通常是中文）
            title_run.font.name = '標楷體'
            title_run._element.rPr.rFonts.set(qn('w:ascii'), '標楷體')
            title_run._element.rPr.rFonts.set(qn('w:hAnsi'), '標楷體')
            title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
            
            # === 內文 ===
            for line in summary_text.split('\n'):
                if line.strip():
                    p = paragraph.insert_paragraph_before(line.strip())
                    p.paragraph_format.line_spacing = 1.5
                    p.paragraph_format.space_after = Pt(6)
                    
                    # 對每一行，可能有多個 run（如果有粗體等）
                    for run in p.runs:
                        run.font.size = Pt(12)
                        
                        # 關鍵：根據文字內容決定字體
                        if any('\u4e00' <= char <= '\u9fff' for char in run.text):  # 有中文
                            run.font.name = '標楷體'
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
                            run._element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
                            run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')
                        else:  # 純英文
                            run.font.name = 'Times New Roman'
                            run._element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
                            run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')  # 保險
            
            bookmark_found = True
            break
    
    # === fallback：書籤沒找到，加到最後 ===
    if not bookmark_found:
        doc.add_page_break()
        
        # 標題
        title_p = doc.add_paragraph("【產生報告總結及建議】")
        title_p.paragraph_format.line_spacing = 1.5
        title_run = title_p.runs[0]
        title_run.bold = True
        title_run.font.size = Pt(16)
        title_run.font.name = '標楷體'
        title_run._element.rPr.rFonts.set(qn('w:ascii'), '標楷體')
        title_run._element.rPr.rFonts.set(qn('w:hAnsi'), '標楷體')
        title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
        
        # 內文
        for line in summary_text.split('\n'):
            if line.strip():
                p = doc.add_paragraph(line.strip())
                p.paragraph_format.line_spacing = 1.5
                p.paragraph_format.space_after = Pt(6)
                
                for run in p.runs:
                    run.font.size = Pt(12)
                    if any('\u4e00' <= char <= '\u9fff' for char in run.text):
                        run.font.name = '標楷體'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
                        run._element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
                        run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')
                    else:
                        run.font.name = 'Times New Roman'
                        run._element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
                        run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
    
    doc.save(output_path or word_path)
    return output_path or word_path

@summary_bp.route('/generate_summary', methods=['POST'])
def generate_summary():
    # 1. 取得參數
    text_input = request.form.get('text_input', '').strip()
    custom_prompt = request.form.get('custom_prompt', '').strip()
    template_type = request.form.get('template_type', 'general')
    
    # 2. 處理檔案上傳（只接受 .docx）
    uploaded_word_path = None
    file = request.files.get('file')
    if file and file.filename.lower().endswith('.docx'):
        filename = secure_filename(file.filename)
        uploaded_word_path = os.path.join(tempfile.gettempdir(), f"upload_{uuid.uuid4()}_{filename}")
        file.save(uploaded_word_path)
    
    # 3. 讀取文字內容（優先 text_input，其次上傳檔案）
    full_content = text_input
    if uploaded_word_path:
        try:
            full_content = read_word_file(uploaded_word_path)  # 你的讀取函數
        except:
            return jsonify({"error": "讀取 Word 檔案失敗"}), 500
    
    if not full_content.strip():
        return jsonify({"error": "請提供文字或上傳 Word 檔案"}), 400
    
    try:
        client, deployment = get_client_and_deployment()
        
        # 選擇模板
        base_prompt = SUMMARY_TEMPLATES.get(template_type, SUMMARY_TEMPLATES["general"])
        
        # 自訂 prompt 優先
        final_prompt = custom_prompt + "\n\n" + base_prompt if custom_prompt else base_prompt
        
        messages = [
            {"role": "system", "content": final_prompt},
            {"role": "user", "content": full_content}
        ]
        
        response = client.chat.completions.create(
            model=deployment,
            messages=messages,
            temperature=0.3,
            max_tokens=1000
        )
        
        summary = response.choices[0].message.content.strip()
        
        # ============ 新增：如果有上傳 Word，插入摘要並回傳檔案 ============
        if uploaded_word_path:
            output_filename = f"AI摘要_{os.path.basename(uploaded_word_path)}"
            output_path = os.path.join(tempfile.gettempdir(), f"output_{uuid.uuid4()}_{output_filename}")
            
            insert_summary_at_bookmark(uploaded_word_path, summary, "AI_SUMMARY_HERE", output_path)
            
            os.remove(uploaded_word_path)
            
            # 修正處：使用 response 物件並加上自定義 Header
            from flask import make_response
            import urllib.parse
            
            response = make_response(send_file(
                output_path,
                as_attachment=True,
                download_name=output_filename,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            ))
            # 將摘要放入 Header (需編碼處理避免中文亂碼)
            response.headers['X-Summary-Text'] = urllib.parse.quote(summary)
            # 允許前端讀取此自訂 Header
            response.headers['Access-Control-Expose-Headers'] = 'X-Summary-Text'
            
            return response
        
    except Exception as e:
        import traceback
        print(traceback.format_exc())
        return jsonify({"error": f"處理失敗：{str(e)}"}), 500